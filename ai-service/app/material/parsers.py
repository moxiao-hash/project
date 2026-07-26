"""受支持资料格式的确定性文本提取器。"""

from io import BytesIO

from docx import Document
from pypdf import PdfReader

from app.material.models import ParsedBlock, ParsedDocument


class UnsupportedMaterialError(ValueError):
    """资料类型不在当前阶段支持范围内。"""


class ScannedPdfError(ValueError):
    """PDF 没有可提取文本层，需要后续 OCR 能力。"""


class MaterialParser:
    """根据 Java 提供的 MaterialType 选择解析方式。"""

    def parse(self, material_type: str, content: bytes) -> ParsedDocument:
        parser = {
            "TEXT": self._parse_text,
            "PASTED_ARTICLE": self._parse_text,
            "MARKDOWN": self._parse_markdown,
            "PDF": self._parse_pdf,
            "WORD": self._parse_docx,
        }.get(material_type)
        if parser is None:
            raise UnsupportedMaterialError(f"不支持的资料类型: {material_type}")
        parsed = parser(content)
        if not parsed.blocks:
            raise ValueError("资料中没有可处理的文字")
        return parsed

    def _parse_text(self, content: bytes) -> ParsedDocument:
        text = content.decode("utf-8-sig").strip()
        return ParsedDocument((ParsedBlock(text=text, locator="正文"),) if text else ())

    def _parse_markdown(self, content: bytes) -> ParsedDocument:
        text = content.decode("utf-8-sig")
        blocks: list[ParsedBlock] = []
        for index, paragraph in enumerate(text.split("\n\n"), start=1):
            cleaned = paragraph.strip()
            if cleaned:
                locator = f"段落 {index}"
                if cleaned.startswith("#"):
                    locator = f"标题 {index}"
                blocks.append(ParsedBlock(text=cleaned, locator=locator))
        return ParsedDocument(tuple(blocks))

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        reader = PdfReader(BytesIO(content))
        blocks = tuple(
            ParsedBlock(text=text, locator=f"第 {page_number} 页")
            for page_number, page in enumerate(reader.pages, start=1)
            if (text := (page.extract_text() or "").strip())
        )
        if not blocks:
            raise ScannedPdfError("PDF 没有可提取文本层，需要 OCR")
        return ParsedDocument(blocks)

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        document = Document(BytesIO(content))
        blocks: list[ParsedBlock] = []
        paragraph_number = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraph_number += 1
                blocks.append(
                    ParsedBlock(text=text, locator=f"段落 {paragraph_number}")
                )
        for table_number, table in enumerate(document.tables, start=1):
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
            ]
            text = "\n".join(row for row in rows if row.strip(" |"))
            if text:
                blocks.append(ParsedBlock(text=text, locator=f"表格 {table_number}"))
        return ParsedDocument(tuple(blocks))
