from io import BytesIO

import pytest

from app.material.parsers import (
    MaterialParser,
    ScannedPdfError,
    UnsupportedMaterialError,
)


def test_decodes_utf8_text_with_source_locator() -> None:
    parsed = MaterialParser().parse("TEXT", "你好，Spring".encode())

    assert parsed.blocks[0].text == "你好，Spring"
    assert parsed.blocks[0].locator == "正文"


def test_extracts_docx_paragraphs_and_tables() -> None:
    from docx import Document

    document = Document()
    document.add_heading("Spring 学习路线", level=1)
    document.add_paragraph("先学习依赖注入。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "知识点"
    table.cell(0, 1).text = "自动配置"
    output = BytesIO()
    document.save(output)

    parsed = MaterialParser().parse("WORD", output.getvalue())

    assert [block.locator for block in parsed.blocks] == ["段落 1", "段落 2", "表格 1"]
    assert "知识点 | 自动配置" in parsed.blocks[-1].text


def test_rejects_pdf_without_text_layer() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    output = BytesIO()
    writer.write(output)

    with pytest.raises(ScannedPdfError, match="OCR"):
        MaterialParser().parse("PDF", output.getvalue())


def test_rejects_legacy_word_document() -> None:
    with pytest.raises(UnsupportedMaterialError):
        MaterialParser().parse("DOC", b"legacy")
